from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOllama
from langchain_community.embeddings import FastEmbedEmbeddings
from langchain.schema.output_parser import StrOutputParser
from langchain_community.document_loaders import PyPDFLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema.runnable import RunnablePassthrough
from langchain.prompts import PromptTemplate
from langchain_community.vectorstores.utils import filter_complex_metadata
import docx
from typing import List, Dict
from langchain.globals import set_verbose, get_verbose

set_verbose(True)  
current_verbose = get_verbose()


class Document:
    def __init__(self, page_content: str, metadata: dict = None):
        self.page_content = page_content
        self.metadata = metadata or {}

class ChatDocument:
    def __init__(self):
        self.vector_store = None
        self.retriever = None
        self.chain = None
        self.model = ChatOllama(model="mistral")
        self.text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=100)
        self.prompt = PromptTemplate.from_template(
            """
            <s> [INST] You are an assistant for question-answering tasks. Use the following pieces of retrieved context 
            to answer the question. If you don't know the answer, just say that you don't know. Use three sentences
             maximum and keep the answer concise. [/INST] </s> 
            [INST] Question: {question} 
            Context: {context} 
            Answer: [/INST]
            """
        )

    def ingest_file(self, file_path: str, file_name: str) -> None:
        """Ingest a file (PDF or DOCX) and prepare it for question answering."""
        if file_name.lower().endswith(".pdf"):
            docs = PyPDFLoader(file_path=file_path).load()
        elif file_name.lower().endswith(".docx"):
            docs = self.load_docx(file_path)
        else:
            raise ValueError("Unsupported file type")

        chunks = self.text_splitter.split_documents(docs)
        chunks = filter_complex_metadata(chunks)

        self.vector_store = Chroma.from_documents(documents=chunks, embedding=FastEmbedEmbeddings())
        self.retriever = self.vector_store.as_retriever(
            search_type="mmr",
            search_kwargs={
                "k": 5,  # Return top 5 results
                "lambda_mult": 0.5,  # Trade-off parameter between relevance and diversity
            },
        )
        self.chain = ({"context": self.retriever, "question": RunnablePassthrough()}
                      | self.prompt
                      | self.model
                      | StrOutputParser())

    def load_docx(self, file_path: str) -> List[Document]:
        """Load and parse a DOCX file into the required format."""
        doc = docx.Document(file_path)
        text = []
        for paragraph in doc.paragraphs:
            text.append(paragraph.text)
        content = '\n'.join(text)
        return [Document(page_content=content, metadata={})]

    def ask_question(self, query: str) -> str:
        """Ask a question to the assistant and get the response."""
        if not self.chain:
            return "Please, add a document first."
        
        response = self.chain.invoke(query)
        return response if isinstance(response, str) else "I couldn't find a good answer. Please try rephrasing your question."

    def clear_data(self) -> None:
        """Clear the current data and reset the state."""
        self.vector_store = None
        self.retriever = None
        self.chain = None
